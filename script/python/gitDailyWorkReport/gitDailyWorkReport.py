import subprocess
import re
from docx import Document
import os
import argparse
import datetime

def get_git_commit_logs(author_name, repo_path=None, start_date=None, end_date=None):
    """获取git提交记录

    Args:
        author_name (_type_): 提交人名称
        repo_path (_type_, optional): 仓库路径. Defaults to None.
        start_date (_type_, optional): 开始时间. Defaults to None.
        end_date (_type_, optional): 截止时间. Defaults to None.

    Returns:
        _type_: 指定日期的git提交记录,默认为当月提交记录。
    """
    if repo_path:
        # 如果提供了仓库路径，切换到该目录下
        repo_path = os.path.normpath(repo_path)  # 将路径规范化为适应当前操作系统的分隔符
        os.chdir(repo_path)

    # 计算开始时间和截止时间 默认为当月时间
    if not start_date:
        current_date = datetime.datetime.now()
        start_date = current_date.replace(day=1).strftime("%Y-%m-%d")
    if not end_date:
        current_date = datetime.datetime.now()
        end_date = (current_date.replace(day=1, month=current_date.month + 1) - datetime.timedelta(days=1)).strftime("%Y-%m-%d")

    print('start_date',start_date)
    print('end_date',end_date)
    
    try:
        git_log = subprocess.check_output(
            ['git', 'log', '--author=' + author_name, '--date=short', f'--since={start_date}', f'--until={end_date}'],
            universal_newlines=True,
            errors='ignore',
            encoding='utf-8'
        )
        return git_log
    except subprocess.CalledProcessError:
        return "Error: Unable to retrieve Git commit logs."

def categorize_commits_by_date(commit_logs):
    """将提交记录按日期分类

    Args:
        commit_logs (_type_): 指定日期的git提交记录

    Returns:
        _type_: 按日期分类的git提交记录
    """
    commits_by_date = {}
    current_date = None

    for line in commit_logs.splitlines():
        # 使用正则表达式提取日期信息
        date_match = re.match(r'Date:\s+(\d{4}-\d{2}-\d{2})', line)
        if date_match:
            current_date = date_match.group(1)
            if current_date not in commits_by_date:
                commits_by_date[current_date] = []

        if current_date:
            if not line.startswith("Date:"):
                commits_by_date[current_date].append(line)

    return commits_by_date

def create_daily_report(commits_by_date):
    """在这里，你可以自定义如何构建每日工作报告

    Args:
        commits_by_date (_type_): 按日期分类的git提交记录

    Returns:
        _type_: 过滤的git提交记录
    """
    report = ""
    for date, commits in commits_by_date.items():
        report += f"Date: {date}\n"
        report += "\n".join(commit for commit in commits if not commit.startswith("commit") and not commit.startswith("Author")) + "\n\n"

    return report

def write_to_word_document(daily_report, author_name):
    """将整理过的git提交记录写入word文档中

    Args:
        daily_report (_type_): 过滤的git提交记录
        author_name (_type_): 提交人名称
    """
    # 获取脚本所在的目录
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 创建Word文档
    doc = Document()
    doc.add_heading("Programmer's Daily Work Report for " + author_name, level=1)

    # 添加每日工作报告到文档
    doc.add_paragraph(daily_report)

    # 构建文档保存路径
    document_path = os.path.join(script_dir, "daily_work_report.docx")

    # 保存文档
    doc.save(document_path)

if __name__ == "__main__":
    # 定义命令行参数
    parser = argparse.ArgumentParser(description="Get Git commit logs for a specific author.")
    parser.add_argument("author_name", type=str, help="Author name to filter Git commit logs.")
    parser.add_argument("--repo_path", type=str, help="Path to the Git repository (optional).")
    parser.add_argument("--start_date", type=str, help="Start date (YYYY-MM-DD) for the time range (optional).")
    parser.add_argument("--end_date", type=str, help="End date (YYYY-MM-DD) for the time range (optional).")

    # 获取命令行参数
    args = parser.parse_args()

    # 指定要获取提交记录的特定提交人的名字
    author_name = args.author_name
    # 获取可选的仓库路径参数
    repo_path = args.repo_path
    # 获取可选的时间段参数
    start_date = args.start_date
    end_date = args.end_date

    git_commit_logs = get_git_commit_logs(author_name, repo_path, start_date, end_date)
    commits_by_date = categorize_commits_by_date(git_commit_logs)
    daily_report = create_daily_report(commits_by_date)

    # 打印每日工作报告
    print("Programmer's Daily Work Report for", author_name)

    # 将工作报告写入Word文档
    write_to_word_document(daily_report, author_name)
